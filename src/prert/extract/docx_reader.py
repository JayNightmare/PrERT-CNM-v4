"""Helpers for extracting normalized text content from DOCX files."""

from __future__ import annotations

from pathlib import Path
import re
from typing import List

from docx import Document  # type: ignore[import-not-found]


_VISIBLE_CLAUSE_RE = re.compile(r"^(?:[1-9]\d*(?:\.[0-9]+){0,4}|A\.[0-9]+(?:\.[0-9]+)?)\b")


def read_docx_text(path: Path) -> str:
    """Read text from a DOCX file as newline-delimited logical paragraphs."""
    doc = Document(str(path))
    lines: List[str] = []
    numbering_state: dict[int, List[int]] = {}

    for paragraph in doc.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            heading_prefix = _heading_number_prefix(paragraph, numbering_state)
            if heading_prefix and not _VISIBLE_CLAUSE_RE.match(text):
                text = f"{heading_prefix} {text}"
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = _clean_text(paragraph.text)
                    if text:
                        lines.append(text)

    return "\n".join(lines)


def _clean_text(value: str) -> str:
    cleaned = value.replace("\u200b", "").replace("\ufeff", "").replace("\u2060", "")
    cleaned = cleaned.replace("\xa0", " ")
    return cleaned.strip()


def _heading_number_prefix(paragraph, numbering_state: dict[int, List[int]]) -> str | None:
    style_name = ""
    if paragraph.style is not None and paragraph.style.name:
        style_name = str(paragraph.style.name)

    if not style_name.lower().startswith("heading"):
        return None

    ppr = paragraph._p.pPr
    if ppr is None or ppr.numPr is None:
        return None
    if ppr.numPr.numId is None or ppr.numPr.ilvl is None:
        return None

    num_id = int(ppr.numPr.numId.val)
    level = int(ppr.numPr.ilvl.val)
    if level < 0:
        return None

    counters = numbering_state.setdefault(num_id, [0] * 9)
    counters[level] += 1
    for idx in range(level + 1, len(counters)):
        counters[idx] = 0

    segments = [str(counters[idx]) for idx in range(level + 1) if counters[idx] > 0]
    if not segments:
        return None
    return ".".join(segments)
